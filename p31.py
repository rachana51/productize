import streamlit as st
from docx import Document
import fitz  # PyMuPDF
import os
import tempfile
import openai

# Set your OpenRouter API key
openai.api_key = "your-openrouter-api-key"
openai.api_base = "https://openrouter.ai/api/v1"  # Optional if using OpenRouter

def extract_text_from_pdf(file):
    pdf_text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            pdf_text += page.get_text()
    return pdf_text

def call_llm_to_generate_data(template_text, report_text):
    prompt = f"""
You are an insurance claims assistant. Use the report below to populate fields in the given template.

### TEMPLATE (FIELDS in brackets):
{template_text}

### PHOTO REPORT:
{report_text}

Return a JSON of field_name: value format for all placeholder fields.
"""
    response = openai.ChatCompletion.create(
        model="mistral-7b",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    return response["choices"][0]["message"]["content"]

def fill_template(template_doc, filled_data):
    for paragraph in template_doc.paragraphs:
        for key, value in filled_data.items():
            placeholder = f"[{key}]"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # Also fill tables
    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in filled_data.items():
                    placeholder = f"[{key}]"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
    return template_doc

st.set_page_config(page_title="GLR Insurance Automation", layout="wide")
st.title("📄 General Loss Report Automation Tool")

template_file = st.file_uploader("Upload Insurance Template (.docx)", type="docx")
pdf_files = st.file_uploader("Upload One or More Photo Reports (.pdf)", type="pdf", accept_multiple_files=True)

if template_file and pdf_files:
    if st.button("Generate Filled Report"):
        with st.spinner("🔍 Extracting text from photo reports..."):
            combined_text = ""
            for pdf in pdf_files:
                combined_text += extract_text_from_pdf(pdf) + "\n"

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_template:
            tmp_template.write(template_file.read())
            tmp_template_path = tmp_template.name

        doc = Document(tmp_template_path)

        # Get raw template text with placeholders
        raw_text = "\n".join([para.text for para in doc.paragraphs])

        with st.spinner("🤖 Calling LLM to extract and match key fields..."):
            llm_response = call_llm_to_generate_data(raw_text, combined_text)
            try:
                filled_data = eval(llm_response)  # Use safer JSON parser in prod
            except Exception as e:
                st.error(f"LLM response error: {e}")
                st.text(llm_response)
                st.stop()

        with st.spinner("📝 Filling in the template..."):
            filled_doc = fill_template(doc, filled_data)

        output_path = os.path.join(tempfile.gettempdir(), "filled_glr_report.docx")
        filled_doc.save(output_path)

        with open(output_path, "rb") as f:
            st.success("✅ Report generated!")
            st.download_button("📥 Download Filled Report", data=f, file_name="Filled_GLReport.docx")

